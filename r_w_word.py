# -*- encoding: utf-8 -*-
"""
@Author      : Alfred T
@Time        : 2020/5/25 18:23
@description : 从文件中读取word然后写入
"""

from win32com import client as wc
import docx
import os


# doc 转 docx 便于后面调用docx库操作
# rootdir 文件夹路径，文件夹下为各word文件
def doc2docx(rootdir):
    word = wc.Dispatch("Word.Application")
    # 得到文件夹下所有的目录与文件名
    list = os.listdir(rootdir)
    # 遍历每个文件
    for i in range(0, len(list)):
        path = os.path.join(rootdir, list[i])
        # 如果文件以.doc结尾
        if path.endswith(".doc"):
            # 将其转换为docx
            doc = word.Documents.Open(path)
            doc.SaveAs(rootdir + list[i] + "x", 12)
            doc.Close()
            # 移除以前的doc文件
            os.remove(path)
    word.Quit()


# 读取文件夹中每个word中的内容
# rootdir:    文件夹的路径
# return:     读取到的内容
def read_word(rootdir):
    list = os.listdir(rootdir)  # 列出文件夹下所有的目录与文件
    # 数组用于存储读取的内容
    texts = []
    for i in range(0, len(list)):
        path = os.path.join(rootdir, list[i])
        # 如果此文件以docx结尾
        if path.endswith(".docx"):
            # 使用docx库打开此word文件
            docStr = docx.Document(path)
            # 遍历文件中的每一段
            for paragraph in docStr.paragraphs:
                parStr = paragraph.text
                # 如果此段内容不为空，将内容放入数组中
                if parStr != "":
                    texts.append("{} : {}".format(list[i][:-5], parStr))
    # 返回数组
    return texts


if __name__ == "__main__":

    rootdir = 'C:/Users/ty/Desktop/宣传改/优秀少先队员/doc/'

    # 转换doc为docx
    doc2docx(rootdir)

    # 获取内容
    texts = read_word(rootdir)

    doc = docx.Document()  # 创建一个Document对象
    # 设置字体
    doc.styles['Normal'].font.name = u'黑体'
    doc.styles['Normal']._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'黑体')

    # 将数组内容全部写入新的文档
    for text in texts:
        doc.add_paragraph(text + "\n")  # 增加一个paragraph
    doc.save(rootdir + '_汇总.docx')
